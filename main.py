import os
from dotenv import load_dotenv
from langchain_community.document_loaders import PyMuPDFLoader

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
# from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_openai import ChatOpenAI
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain.chains import create_retrieval_chain
from langchain.schema.runnable import RunnableParallel, RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser

from docx import Document as DocxDocument  # Rename to avoid conflicts
from langchain.schema import Document

from langchain.prompts import ChatPromptTemplate


load_dotenv()

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("Missing Google API Key. Please check your .env file.")

OPENAI_API_KEY = os.getenv("OPENAI_KEY")
if not OPENAI_API_KEY:
    raise ValueError("Missing OpenAI API Key. Please check your .env file.")



DOC_FOLDER = "data/pdf_files"
CHROMA_PATH = "vector_db_pdf"
MODEL_NAME = "gpt-4o-mini"
TEMPERATURE = 0.7
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200


def load_docx_documents(folder_path):
    docx_docs = []
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            try:
                doc = DocxDocument(file_path)
                doc_content = "\n".join([para.text for para in doc.paragraphs])
                docx_docs.append(Document(page_content=doc_content, metadata={"source": filename}))
            except Exception as e:
                print(f"Error loading {filename}: {e}")
    return docx_docs

docx_docs = load_docx_documents(DOC_FOLDER)
if not docx_docs or not all(isinstance(doc, Document) for doc in docx_docs):
    raise ValueError("No valid LangChain Document objects found in the folder.")


splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP,
    separators=["\n\n", "\n", " ", ""]
)
chunks = splitter.split_documents(docx_docs)
print(f"Total chunks created: {len(chunks)}")




# embedding_model = OpenAIEmbeddings()
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
vector_db = Chroma.from_documents(
    chunks, embeddings, persist_directory=CHROMA_PATH
)
print("Vector DB created successfully.")

retriever = vector_db.as_retriever(search_type="mmr", search_kwargs={"k": 3})

def query_documents(query):
    retrieved_docs = retriever.invoke(query)
    final_docs = retrieved_docs[:3]
    context = "\n".join([doc.page_content for doc in final_docs])[:10000]
    return context



SYSTEM_PROMPT = """
คุณเป็นแชทบอทอัจฉริยะที่เชี่ยวชาญเรื่องหนังสือเตรียมสอบข้าราชการ 🎓 
- ให้ข้อมูลเกี่ยวกับหนังสือ รวมถึงเนื้อหา จุดเด่น และวิธีใช้  
- ตอบคำถามเกี่ยวกับการสอบข้าราชการ เช่น โครงสร้างข้อสอบ แนวข้อสอบ และเทคนิคการเตรียมตัว  
- แนะนำสินค้าที่เหมาะสมตามความต้องการของลูกค้า  
- ตอบคำถามอย่างสุภาพ ชัดเจน และเป็นมิตร โดยใช้ **ภาษาไทย** ที่เป็นธรรมชาติ  
- หลีกเลี่ยงการให้ข้อมูลที่ไม่มีในหนังสือ หรือไม่สามารถยืนยันได้  
- กระตุ้นให้ลูกค้าซื้อหนังสือโดยแนะนำ **จุดเด่นของหนังสือ** และ **ประโยชน์ที่ได้รับ**
- ตอนสั้น ๆ กระชับ ได้ใจความ เหมือนคุยกับคน
"""


chat_prompt = ChatPromptTemplate.from_messages([
    ("system", SYSTEM_PROMPT),
    ("human", "{query}")  # Placeholder for customer questions
])


llm = ChatOpenAI(
    model_name=MODEL_NAME, 
    temperature=TEMPERATURE, 
    api_key=OPENAI_API_KEY,
    )

retrieval_chain = (
    RunnableParallel(
        {"context": retriever, "query": RunnablePassthrough()}
    )
    | chat_prompt
    | llm
    | StrOutputParser()
)



# llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro", temperature=TEMPERATURE)
# qa_chain = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever)


query = "มีหนังสือเกียวกับภาษาไทยไหม"
context = query_documents(query)
print("Context:\n", context)
response = retrieval_chain.invoke(f"{query}\nข้อมูลที่เกี่ยวข้อง:\n{context}")
# response = retrieval_chain.invoke(query)

print("Response:\n", response)